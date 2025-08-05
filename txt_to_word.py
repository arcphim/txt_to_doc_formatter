#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将txt或md文档转换为格式化的Word文档
"""

import sys
import os
import re
import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from doc_formatter import DocFormatter


def read_txt_file(file_path):
    """读取txt文件内容"""
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()


def read_md_file(file_path):
    """读取md文件内容"""
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()


def parse_txt_content(content):
    """解析txt内容为段落列表"""
    # 按行分割内容
    lines = content.split('\n')
    paragraphs = []
    current_paragraph = ""
    
    for line in lines:
        # 如果是空行，表示一个段落结束
        if not line.strip():
            if current_paragraph:
                paragraphs.append(current_paragraph.strip())
                current_paragraph = ""
        else:
            # 继续构建当前段落
            if current_paragraph:
                current_paragraph += "\n" + line
            else:
                current_paragraph = line
    
    # 添加最后一个段落
    if current_paragraph:
        paragraphs.append(current_paragraph.strip())
    
    return paragraphs


def parse_md_content(content):
    """解析md内容为段落列表"""
    # 按行分割内容
    lines = content.split('\n')
    paragraphs = []
    current_paragraph = ""
    
    # Markdown标题标识
    md_headings = [r'^#{1,6}\s', r'^[=\-]+$']
    
    for line in lines:
        # 检查是否为Markdown标题
        is_heading = any(re.match(pattern, line.strip()) for pattern in md_headings)
        
        # 如果是空行或标题，表示一个段落结束
        if not line.strip() or is_heading:
            if current_paragraph:
                paragraphs.append(current_paragraph.strip())
                current_paragraph = ""
            # 如果是标题，单独作为一个段落，并去除#符号
            if is_heading:
                # 去除标题行中的#符号
                heading_text = line.strip()
                # 使用正则表达式去除开头的#符号和空格
                heading_text = re.sub(r'^#{1,6}\s*', '', heading_text)
                paragraphs.append(heading_text)
        else:
            # 继续构建当前段落
            if current_paragraph:
                current_paragraph += " " + line
            else:
                current_paragraph = line
    
    # 添加最后一个段落
    if current_paragraph:
        paragraphs.append(current_paragraph.strip())
    
    return paragraphs


def create_word_document(paragraphs, output_path, config_path='config.json', is_markdown=False):
    """创建格式化的Word文档"""
    # 创建DocFormatter实例
    formatter = DocFormatter(config_path)
    
    # 创建新的Word文档
    doc = Document()
    
    # 添加段落到文档
    for paragraph_text in paragraphs:
        # 添加段落
        para = doc.add_paragraph()
        para.add_run(paragraph_text)
    
    # 应用格式化
    # 设置页边距
    formatter.set_page_margins(doc)
    
    # 设置标题格式（假设第一个段落是标题）
    if len(doc.paragraphs) > 0:
        formatter.format_title(doc)
    
    # 设置标题层级
    if is_markdown:
        # 处理Markdown格式的标题
        processed_paragraphs = format_markdown_headings(doc)
    else:
        # 处理中文格式的标题
        processed_paragraphs = formatter.format_headings(doc)
    
    # 设置正文格式
    formatter.format_body(doc, processed_paragraphs)
    
    # 设置页码
    formatter.set_page_numbers(doc)
    
    # 保存文档
    doc.save(output_path)
    print(f"文档已保存为: {output_path}")


def format_markdown_headings(doc):
    """设置Markdown标题层级格式"""
    # 用于标记已处理的段落
    processed_paragraphs = set()
    
    # 加载配置
    with open('config.json', 'r', encoding='utf-8') as f:
        config = json.load(f)
    
    # 处理文章标题（第一个段落）
    if len(doc.paragraphs) > 0:
        # 创建DocFormatter实例来使用format_title方法
        formatter = DocFormatter('config.json')
        formatter.format_title(doc)
        processed_paragraphs.add(0)
    
    # 定义Markdown标题的正则表达式（从二级标题开始）
    heading_patterns = [
        r'^##\s+(.*)$',     # 二级标题
        r'^###\s+(.*)$',    # 三级标题
        r'^####\s+(.*)$',   # 四级标题
        r'^#####\s+(.*)$',  # 五级标题
        r'^######\s+(.*)$'  # 六级标题
    ]
    
    heading_configs = config['heading_levels']
    
    # 从第二个段落开始处理（跳过文章标题）
    for i, para in enumerate(doc.paragraphs[1:], 1):
        text = para.text.strip()
        
        # 检查是否为Markdown标题
        for level, pattern in enumerate(heading_patterns):
            # level+1是因为我们跳过了第一个标题级别
            match = re.match(pattern, text)
            if match and (level+1) < len(heading_configs):
                # 提取标题内容
                heading_text = match.group(1)
                # 清空段落
                para.clear()
                # 添加标题内容
                run = para.add_run(heading_text)
                
                # 应用相应的格式
                # level+1是因为我们跳过了第一个标题级别
                heading_config = heading_configs[level+1]
                run.font.name = heading_config['font']
                run._element.rPr.rFonts.set(qn('w:eastAsia'), heading_config['font'])
                run.font.size = Pt(heading_config['size'])
                if 'bold' in heading_config and heading_config['bold']:
                    run.bold = True
                
                # 设置段落格式
                para.paragraph_format.first_line_indent = Pt(14.4 * heading_config.get('indent', 2))
                para.paragraph_format.line_spacing_rule = 4  # 固定值
                para.paragraph_format.line_spacing = Pt(config['spacing']['line_spacing'])
                
                # 标记为已处理
                processed_paragraphs.add(i)
                break
    
    return processed_paragraphs

def txt_to_word(input_path, output_path, config_path='config.json'):
    """将txt文件转换为格式化的Word文档"""
    print(f"开始处理文件: {input_path}")
    # 检查文件扩展名
    _, ext = os.path.splitext(input_path)
    print(f"文件扩展名: {ext}")
    
    if ext.lower() == '.txt':
        print("读取txt文件...")
        # 读取txt文件
        content = read_txt_file(input_path)
        print("解析txt内容...")
        # 解析内容
        paragraphs = parse_txt_content(content)
        is_markdown = False
    elif ext.lower() == '.md':
        print("读取md文件...")
        # 读取md文件
        content = read_md_file(input_path)
        print("解析md内容...")
        # 解析内容
        paragraphs = parse_md_content(content)
        is_markdown = True
    else:
        raise ValueError(f"不支持的文件格式: {ext}")
    
    print("创建Word文档...")
    # 创建Word文档
    create_word_document(paragraphs, output_path, config_path, is_markdown)
    print("Word文档创建完成")


def main():
    if len(sys.argv) < 3:
        print("使用方法: python txt_to_word.py <输入文件> <输出文件> [配置文件]")
        print("示例: python txt_to_word.py input.txt output.docx")
        sys.exit(1)
    
    input_path = sys.argv[1]
    output_path = sys.argv[2]
    config_path = sys.argv[3] if len(sys.argv) > 3 else 'config.json'
    
    # 检查输入文件是否存在
    if not os.path.exists(input_path):
        print(f"错误: 输入文件不存在: {input_path}")
        sys.exit(1)
    
    try:
        txt_to_word(input_path, output_path, config_path)
        print("转换完成!")
    except Exception as e:
        print(f"转换过程中出错: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()