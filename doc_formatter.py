import json
import os
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
import re
from docx.oxml.ns import qn
from docx.shared import RGBColor
import re
import logging

class DocFormatter:
    def __init__(self, config_path='config.json'):
        self.config = self.load_config(config_path)
        
    def load_config(self, config_path):
        """加载配置文件"""
        with open(config_path, 'r', encoding='utf-8') as f:
            return json.load(f)
        
    def format_title(self, doc):
        """设置标题格式"""
        if len(doc.paragraphs) == 0:
            return
        
        title_config = self.config['title_font']
        title_paragraph = doc.paragraphs[0]
        original_text = title_paragraph.text
        title_paragraph.clear()
        
        # 获取数字字体设置
        digit_font = self.config['body_font']['digit_font']
        
        # 分离英文/数字和汉字
        parts = re.split(r'([a-zA-Z0-9]+)', original_text)
        for part in parts:
            if part.strip() == '':
                continue
            run = title_paragraph.add_run(part)
            if re.match(r'^[a-zA-Z0-9]+$', part):
                # 英文或数字部分
                run.font.name = digit_font
                run._element.rPr.rFonts.set(qn('w:eastAsia'), digit_font)
            else:
                # 汉字部分
                run.font.name = title_config['name']
                run._element.rPr.rFonts.set(qn('w:eastAsia'), title_config['name'])
            run.font.size = Pt(title_config['size'])
            run.bold = title_config.get('bold', False)
        
        # 设置居中对齐
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    def format_headings(self, doc):
        """设置标题层级格式"""
        heading_configs = self.config['heading_levels']
        
        # 用于标记已处理的段落
        processed_paragraphs = set()
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            
            # 检查是否为一级标题 (支持一、二、三等)
            if re.match(r'^[一二三四五六七八九十]、.*', text):
                self.apply_heading_style(para, heading_configs[0])
                processed_paragraphs.add(i)
            # 检查是否为二级标题 (支持（一）（二）（三）等)
            elif re.match(r'^（[一二三四五六七八九十]+）.*', text):
                self.apply_heading_style(para, heading_configs[1])
                processed_paragraphs.add(i)
            # 检查是否为三级标题 (支持1. 2. 3. 等)
            elif re.match(r'^\d+\. .*', text):
                self.apply_heading_style(para, heading_configs[2])
                processed_paragraphs.add(i)
            # 检查是否为四级标题 (支持（1）（2）（3）等)
            elif re.match(r'^（\d+）.*', text):
                self.apply_heading_style(para, heading_configs[3])
                processed_paragraphs.add(i)
        
        return processed_paragraphs
    
    def format_body(self, doc, processed_paragraphs=None):
        """设置正文格式"""
        body_config = self.config['body_font']
        line_spacing = self.config['spacing']['line_spacing']
        digit_font = self.config['body_font']['digit_font']
        
        # 如果没有提供已处理的段落索引，则处理所有正文段落
        if processed_paragraphs is None:
            # 跳过第一个段落（主标题）
            paragraphs_to_process = enumerate(doc.paragraphs[1:], 1)
        else:
            # 只处理未被标记为标题的段落
            paragraphs_to_process = [(i, para) for i, para in enumerate(doc.paragraphs[1:], 1) if i not in processed_paragraphs]
        
        for i, para in paragraphs_to_process:
            # 跳过空段落或只包含空白字符的段落
            if not para.text.strip():
                continue
                
            # 设置段落行距为固定值28磅
            para.paragraph_format.line_spacing_rule = 4  # WD_LINE_SPACING.FIXED
            para.paragraph_format.line_spacing = Pt(line_spacing)
            
            # 设置正文段落对齐方式为左对齐
            para.alignment = 0  # WD_ALIGN_PARAGRAPH.LEFT
            
            # 设置首行缩进为2个汉字宽度（与标题格式保持一致）
            para.paragraph_format.first_line_indent = Pt(14.4 * 2)
            
            # 处理正文字体和数字字体，分离英文/数字和汉字
            original_text = para.text
            para.clear()
            
            # 分离英文/数字和汉字，包括常用符号
            parts = re.split(r'([a-zA-Z0-9%$€¥℃℉°+\-×÷‰‱]+)', original_text)
            for part in parts:
                if part.strip() == '':
                    continue
                run = para.add_run(part)
                if re.match(r'^[a-zA-Z0-9%$€¥℃℉°+\-×÷‰‱]+$', part):
                    # 英文或数字部分
                    run.font.name = digit_font
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), digit_font)
                else:
                    # 汉字部分
                    run.font.name = body_config['name']
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), body_config['name'])
                run.font.size = Pt(body_config['size'])
        
    def apply_heading_style(self, para, config):
        """应用标题样式"""
        # 保存原始文本
        original_text = para.text
        para.clear()
        
        # 获取数字字体设置
        digit_font = self.config['body_font']['digit_font']
        
        # 处理3级和4级标题中的数字
        if config.get('level') in [3, 4]:
            if config['level'] == 3:
                # 匹配三级标题格式 (1. )
                match = re.match(r'^(\d+\. )(.*)$', original_text)
            else:
                # 匹配四级标题格式 (（1）)
                match = re.match(r'^（(\d+)）(.*)$', original_text)
            
            if match:
                if config['level'] == 3:
                    # 三级标题直接使用捕获的前缀（包含点号）
                    prefix = match.group(1)
                else:
                    # 四级标题需要重新构造前缀（包含括号）
                    prefix = f"（{match.group(1)}）"
                content = match.group(2)
                
                # 设置数字部分字体为Times New Roman
                prefix_run = para.add_run(prefix)
                prefix_run.font.name = digit_font
                prefix_run._element.rPr.rFonts.set(qn('w:eastAsia'), digit_font)
                prefix_run.font.size = Pt(config['size'])
                prefix_run.bold = config['bold']
                
                # 设置标题内容字体，分离英文/数字和汉字，包括常用符号
                parts = re.split(r'([a-zA-Z0-9%$€¥℃℉°+\-×÷‰‱]+)', content)
                for part in parts:
                    if part.strip() == '':
                        continue
                    content_run = para.add_run(part)
                    if re.match(r'^[a-zA-Z0-9%$€¥℃℉°+\-×÷‰‱]+$', part):
                        # 英文或数字部分
                        content_run.font.name = digit_font
                        content_run._element.rPr.rFonts.set(qn('w:eastAsia'), digit_font)
                    else:
                        # 汉字部分
                        content_run.font.name = config['font']
                        content_run._element.rPr.rFonts.set(qn('w:eastAsia'), config['font'])
                    content_run.font.size = Pt(config['size'])
                    content_run.bold = config['bold']
            else:
                # 未匹配到时使用默认样式
                run = para.add_run(original_text)
                run.font.name = config['font']
                run._element.rPr.rFonts.set(qn('w:eastAsia'), config['font'])
                run.font.size = Pt(config['size'])
                run.bold = config['bold']
        else:
            # 1-2级标题处理，分离英文/数字和汉字，包括常用符号
            parts = re.split(r'([a-zA-Z0-9%$€¥℃℉°+\-×÷‰‱]+)', original_text)
            for part in parts:
                if part.strip() == '':
                    continue
                run = para.add_run(part)
                if re.match(r'^[a-zA-Z0-9%$€¥℃℉°+\-×÷‰‱]+$', part):
                    # 英文或数字部分
                    run.font.name = digit_font
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), digit_font)
                else:
                    # 汉字部分
                    run.font.name = config['font']
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), config['font'])
                run.font.size = Pt(config['size'])
                run.bold = config['bold']
        
        # 设置首字前空2格
        para.paragraph_format.first_line_indent = Pt(14.4 * 2)  # 2个汉字宽度
        
        # 设置段落行距
        para.paragraph_format.line_spacing_rule = 4  # 固定值
        para.paragraph_format.line_spacing = Pt(self.config['spacing']['line_spacing'])
        
    def set_page_margins(self, doc):
        """设置页边距"""
        margins = self.config['margins']
        section = doc.sections[0]
        
        section.top_margin = Cm(margins['top'])
        section.bottom_margin = Cm(margins['bottom'])
        section.left_margin = Cm(margins['left'])
        section.right_margin = Cm(margins['right'])
        
    def set_page_numbers(self, doc):
        """设置页码"""
        page_config = self.config['page_number']
        section = doc.sections[0]
        
        # 添加页码
        footer = section.footer
        para = footer.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加页码字段
        run = para.add_run()
        # 创建页码字段元素
        fld_char = OxmlElement('w:fldChar')
        fld_char.set(qn('w:fldCharType'), 'begin')
        run._element.append(fld_char)
        
        instr_text = OxmlElement('w:instrText')
        instr_text.text = 'PAGE'
        run._element.append(instr_text)
        
        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')
        run._element.append(fld_char_end)
        
    logging.basicConfig(
        filename='app.log',
        level=logging.DEBUG,
        format='%(asctime)s - %(levelname)s - %(message)s',
        encoding='utf-8'  # 添加编码配置
    )
    
    def format_document(self, input_path, output_path):
        # 检查输入文件扩展名
        _, ext = os.path.splitext(input_path)
        
        if ext.lower() in ['.txt', '.md']:
            # 对于txt/md文件，应该使用txt_to_word模块处理，而不是在这里处理
            raise ValueError(f"format_document方法不支持处理{ext}文件，请使用txt_to_word模块处理")
        
        try:
            logging.debug(f"尝试打开输入文件: {input_path}")
            doc = Document(input_path)
            
            # 设置页边距
            self.set_page_margins(doc)
            
            # 设置标题格式
            self.format_title(doc)
            
            # 设置标题层级
            processed_paragraphs = self.format_headings(doc)
            
            # 设置正文格式
            self.format_body(doc, processed_paragraphs)
            
            # 设置页码
            self.set_page_numbers(doc)
            
            # 保存文档
            logging.debug(f"尝试保存输出文件: {output_path}")
            doc.save(output_path)
            logging.info(f"文件保存成功: {output_path}")
            return output_path
        except PermissionError as e:
            logging.error(f"权限错误: {str(e)}, 文件路径: {output_path}")
            raise
        except Exception as e:
            logging.error(f"处理文档时出错: {str(e)}")
            raise